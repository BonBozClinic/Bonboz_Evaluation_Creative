# -*- coding: utf-8 -*-
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BD = Path(__file__).parent
data = json.load(open(BD/"results_b2.json", encoding="utf-8"))
data.sort(key=lambda x: int(x["id"]))

MINT="2E8B7A"; DARK="1F3B35"; GREY="666666"; BOX="EAF5F1"
REDc="F8D7DA"; GREENc="D9EFDC"; AMBERc="FFF1CC"; REDt="C0392B"; GREENt="2E7D32"; AMBERt="B8860B"

doc=Document(); sec=doc.sections[0]; sec.page_height=Mm(297); sec.page_width=Mm(210)
sec.top_margin=sec.bottom_margin=Mm(14); sec.left_margin=sec.right_margin=Mm(15)
st=doc.styles["Normal"]; st.font.name="Arial"; st.font.size=Pt(10); st.element.rPr.rFonts.set(qn("w:eastAsia"),"Arial")

def shade(el,fill):
    pr=el.get_or_add_tcPr() if el.tag.endswith('}tc') else el.get_or_add_pPr()
    sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),fill); pr.append(sh)
def run(p,t,b=False,it=False,col=None,sz=None,font=None):
    r=p.add_run(t); r.bold=b; r.italic=it
    if col: r.font.color.rgb=RGBColor.from_string(col)
    if sz: r.font.size=Pt(sz)
    if font: r.font.name=font
    return r
def H(t,sz=13,col=MINT,before=12,after=4):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); run(p,t,b=True,col=col,sz=sz); return p
def sub(t,col=DARK):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2); run(p,t,b=True,col=col,sz=10.5); return p
def bullets(items,sz=9.5,maxn=None):
    items=items[:maxn] if maxn else items
    for x in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Pt(12); p.paragraph_format.space_after=Pt(1); run(p,x,sz=sz)

VC={"SẴN SÀNG CHẠY":(GREENc,GREENt),"SỬA POLICY TRƯỚC":(AMBERc,AMBERt),"TỐI ƯU CHẤT LƯỢNG TRƯỚC":(AMBERc,AMBERt),"KHÔNG CHẠY":(REDc,REDt)}
QC={"ĐẠT":(GREENc,GREENt),"CẦN TỐI ƯU":(AMBERc,AMBERt),"YẾU":(REDc,REDt)}
DIM=[("vd_hook","Hook"),("vd_hold","Hold"),("vd_message","Msg"),("vd_cta","CTA"),("vd_audience","Aud"),("vd_visual","Visual"),("vd_pacing","Pacing"),("vd_emotion","Emot"),("vd_trust","Trust")]

# TITLE
p=doc.add_paragraph(); run(p,"ĐÁNH GIÁ CREATIVE — Flow Gemini + Claude (A/B/C)",b=True,col=DARK,sz=18)
run(doc.add_paragraph(),"Bonboz Clinic · %d video · Gemini xem+nghe → Claude A∥B verify → C merge · perception + đánh giá đầy đủ"%len(data),col=GREY,sz=9.5)

# OVERVIEW
H("Tổng quan")
hdr=["#","Creative","Điểm","Chất lượng","Policy","Quyết định cuối"]; cw=[Mm(10),Mm(58),Mm(15),Mm(28),Mm(24),Mm(45)]
t=doc.add_table(rows=1,cols=6); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
for i,h in enumerate(hdr):
    c=t.rows[0].cells[i]; c.width=cw[i]; shade(c._tc,MINT); pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.LEFT if i==1 else WD_ALIGN_PARAGRAPH.CENTER; run(pp,h,b=True,col="FFFFFF",sz=9)
for d in data:
    sc=d["scorecard"]; rc=t.add_row().cells
    vals=[d["id"], d["title"], str(sc["total"]), sc["quality"], sc["policy"], sc["final"]]
    for i,val in enumerate(vals):
        rc[i].width=cw[i]; pp=rc[i].paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.LEFT if i==1 else WD_ALIGN_PARAGRAPH.CENTER
        col=None;bold=False
        if i==3 and val in QC: f,col=QC[val]; shade(rc[i]._tc,f); bold=True
        if i==5 and val in VC: f,col=VC[val]; shade(rc[i]._tc,f); bold=True
        if i in(0,2):bold=True
        run(pp,str(val),b=bold,col=col,sz=8.8)

# PER VIDEO
for d in data:
    sc=d["scorecard"]; per=d.get("perception",{}); fcol=VC.get(sc["final"],(BOX,AMBERt))[1]
    hp=doc.add_paragraph(); shade(hp._p,MINT); hp.paragraph_format.space_before=Pt(14); hp.paragraph_format.space_after=Pt(0); run(hp,"#%s · %s"%(d["id"],d["title"]),b=True,col="FFFFFF",sz=12)
    vp=doc.add_paragraph(); run(vp,"Quyết định: ",b=True,sz=10); run(vp,sc["final"],b=True,col=fcol,sz=11); run(vp,"   · Điểm %s · CL %s · Policy %s%s"%(sc["total"],sc["quality"],sc["policy"]," · ⚠ cần người duyệt" if d.get("escalate") else ""),sz=9.5)

    # Claude eval — dim table
    sub("🧠 Claude đánh giá")
    rs=sc.get("dim_reasons",{}); sco=sc.get("scores",{})
    for did,lbl in DIM:
        if did not in sco: continue
        v=round(float(sco[did]),1); bc=REDt if v<5 else (GREENt if v>=7 else MINT)
        pp=doc.add_paragraph(); pp.paragraph_format.space_after=Pt(0); run(pp,"%s: "%lbl,b=True,sz=9.3); run(pp,str(v),b=True,col=bc,sz=9.3)
        if rs.get(did): run(pp," — "+str(rs[did]),sz=9.3,col=GREY)
    # findings grouped
    fs=d.get("findings",[])
    def grp(title,color,filt):
        g=[x for x in fs if filt(x)]
        if not g: return
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3); run(p,title+(" (%d)"%len(g)),b=True,col=color,sz=9.5)
        bullets(["%s — %s"%(("%.0f:%02d"%(int(x["time"]//60),int(x["time"]%60))), x["text"]) for x in sorted(g,key=lambda y:y["time"])],sz=9.2)
    grp("Phải sửa",REDt,lambda x:not x.get("unverified") and x.get("severity") in("CAO","TB"))
    grp("Nên cải thiện",AMBERt,lambda x:not x.get("unverified") and x.get("severity")=="NHẸ")
    grp("Chưa chắc — cần xem",GREY,lambda x:x.get("unverified"))

    # Gemini perception
    sub("👁👂 Gemini quan sát (perception)")
    if per.get("hook_0_3s"):
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(1);run(p,"Hook 0-3s: ",b=True,sz=9.3);run(p,str(per["hook_0_3s"]),sz=9.3)
    if per.get("pacing"):
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(1);run(p,"Pacing: ",b=True,sz=9.3);run(p,str(per["pacing"]),sz=9.3)
    if per.get("audio_profile"):
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(1);run(p,"Audio: ",b=True,sz=9.3);run(p,str(per["audio_profile"]),sz=9.3)
    tr=per.get("transcript",[])
    if tr:
        run(doc.add_paragraph(),"Transcript (voice):",b=True,sz=9.3)
        bullets(["[%.1fs] %s"%(x.get("t",0),x.get("text","")) for x in tr],sz=9)
    ov=per.get("onscreen_text",[])
    if ov:
        run(doc.add_paragraph(),"Text overlay:",b=True,sz=9.3)
        bullets(["[%.1fs] %s"%(x.get("t",0),x.get("text","")) for x in ov],sz=9)
    sh=per.get("shots",[])
    if sh:
        run(doc.add_paragraph(),"Cảnh động (shots):",b=True,sz=9.3)
        bullets(["[%.1fs] %s"%(x.get("t",0),x.get("desc","")) for x in sh],sz=9)
    cl=per.get("claims_verbatim",[])
    if cl:
        run(doc.add_paragraph(),"Claim nguyên văn:",b=True,sz=9.3)
        bullets([str(x) for x in cl],sz=9)

out=r"C:\Users\ADMIN\Downloads\Bonboz_DanhGia_8Video_Flow.docx"
doc.save(out); print("OK ->",out,"| videos:",len(data))
